# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document()

# Title
title = doc.add_heading('南瓜视频号 4月运营分析报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('数据周期：2026年4月1日 - 4月30日')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 1
doc.add_heading('一、每日用户活跃与播放量趋势', level=1)

p1 = doc.add_paragraph(
    '4月期间，南瓜视频号整体呈现"脉冲式"增长特征，单日播放量波动剧烈，'
    '峰值出现在4月18日（播放17.3万），远超月均值1.7万。'
    '用户互动指标（推荐/喜欢/分享）与播放量高度同步，说明内容热度与用户参与度强相关。'
    '新增关注者在爆款日（4/18）达到峰值193人，为日常的5-10倍。'
    '整体来看，4月总播放量51.5万，新增关注489人，但分布极不均衡，'
    '绝大部分流量集中于少数爆款内容。'
)

doc.add_picture(
    r'D:\mazal-folder\projects\data-analysis\南瓜_4月_每日趋势图.png',
    width=Inches(6.5)
)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 2
doc.add_heading('二、视频号与剧目收益/播放量占比', level=1)

p2 = doc.add_paragraph(
    '从账号维度看，广告收益高度集中：南瓜动漫短剧场（63.3%）与飞姐漫剧推荐（29.9%）'
    '两账号合计贡献93%的收益，但播放量仅占总量的76%，说明两者变现效率存在差异。'
    '从剧目维度看，新婚后一剧独占53.2%的收益，而金牌替身播放量最高（50万）'
    '却仅贡献19.6%收益，两者形成鲜明对比——高流量不等于高变现。'
    '值得注意的是，部分小播放量账号/剧目（如天上飞的棉花糖呐、渊狐记）'
    '反而实现较高的收益转化，说明内容类型与变现路径的匹配度比单纯流量更重要。'
)

doc.add_picture(
    r'D:\mazal-folder\projects\data-analysis\南瓜_4月_占比分析_四张饼图.png',
    width=Inches(6.5)
)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 3 - Key Stats Table
doc.add_heading('三、关键数据摘要', level=1)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '指标'
hdr_cells[1].text = '数值'
hdr_cells[0].paragraphs[0].runs[0].bold = True
hdr_cells[1].paragraphs[0].runs[0].bold = True

stats = [
    ('4月总播放量', '515,170'),
    ('4月总广告收益', '1,387.55 元'),
    ('月均日播放量', '17,172'),
    ('单日播放峰值', '173,096（4月18日）'),
    ('峰值互动（喜欢/分享/关注）', '678 / 492 / 193'),
    ('4月新增关注', '489'),
    ('Top1 账号收益占比', '南瓜动漫短剧场 63.3%'),
    ('Top1 剧目收益占比', '新婚后，大叔全家爆宠我 53.2%'),
]

for k, v in stats:
    row_cells = table.add_row().cells
    row_cells[0].text = k
    row_cells[1].text = v

doc.add_paragraph()

# Save
output_path = r'D:\mazal-folder\projects\data-analysis\南瓜_4月运营分析报告.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
