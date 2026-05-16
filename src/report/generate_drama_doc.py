# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document()

# Title
title = doc.add_heading('剧目分类标签表', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('数据来源：南瓜-视频号4月（2026年4月）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Classification mapping
# Format: (剧名, 描述, 频道, 题材)
drama_data = [
    ('新婚后，大叔全家爆宠我', '婚后大叔全家爆宠我，动态漫', '女频', '都市'),
    ('金牌替身', '金牌替身，热门动态漫', '女频', '都市'),
    ('替嫁后，沉睡老公超会撩', '替嫁后沉睡老公超会撩，热门动态漫抢先看', '女频', '都市'),
    ('家有财神小福星日进斗金', '家有财神小福星日进斗金，AI真人解说漫', '女频', '都市'),
    ('渊狐记：重生不负心', '渊狐记：重生不负心，AI漫剧', '男频', '玄幻'),
    ('极热65度：老公中奖2亿当天我连夜回村独美', '极热65度，AI真人解说漫', '女频', '都市'),
    ('双鲤换生，这一世我所向披靡', '双鲤换生，这一世我所向披靡，AI真人解说漫', '男频', '玄幻'),
    ('入住月子中心当天老公让我写借条', '入住月子中心当天老公让我写借条，AI漫剧', '女频', '都市'),
    ('盲眼妖师', '盲眼妖师', '男频', '玄幻'),
    ('战神后妈掀桌了', '战神后妈掀桌了，一口气免费看完', '女频', '都市'),
    ('乙游系统觉醒，我在星际撩疯了', '乙游系统觉醒，我在星际撩疯了，热门动态漫', '女频', '都市'),
    ('我在诡异游戏当食堂阿姨', '我在诡异游戏当食堂阿姨，AI漫剧', '女频', '末世'),
    ('祭品公主的生存法则', '祭品公主的生存法则', '女频', '玄幻'),
    ('9个南瓜换我余生自由', '9个南瓜换我余生自由', '女频', '都市'),
    ('末世冥婚，重生后我驭诡称王', '末世冥婚第二集，热门漫剧', '男频', '末世'),
    ('全城饿疯我靠养猪救苍生', '全城饿疯我靠养猪救苍生', '男频', '末世'),
    ('重生回挑选异兽那天', '重生回挑选异兽那天，AI漫剧', '男频', '玄幻'),
    ('诡异游戏boss竟是我男友', '诡异游戏boss竟是我男友，AI漫剧', '女频', '末世'),
    ('天生贵命', '天生贵命', '女频', '都市'),
    ('诡行异闻录之怒水惊魂', '古装悬疑剧，堪比电视剧的漫剧', '男频', '玄幻'),
    ('女儿失踪后，警察给我带回了两个', '女儿失踪后警察给我带回了两个，AI漫剧', '女频', '都市'),
    ('奶娃魔尊', '魔尊转世、修仙爽文、萌娃逆袭', '男频', '玄幻'),
    ('帝凰之神医弃妃', '帝凰之神医弃妃，动态漫', '女频', '玄幻'),
    ('夏可可的末世求生日记', '夏可可的末世求生日记', '女频', '末世'),
    ('超绝虎王', '超绝虎王', '男频', '玄幻'),
    ('农女殊色', '农女殊色', '女频', '都市'),
    ('自爆祖师', '自爆祖师', '男频', '玄幻'),
]

# ========== Table ==========
doc.add_heading('剧目标签明细', level=1)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '剧名'
hdr_cells[1].text = '描述'
hdr_cells[2].text = '频道'
hdr_cells[3].text = '题材'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

# Data rows
for name, desc, channel, genre in drama_data:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = desc
    row_cells[2].text = channel
    row_cells[3].text = genre

doc.add_paragraph()

# ========== Stats Summary ==========
doc.add_heading('分类统计', level=1)

# Count by channel
channel_count = {}
genre_count = {}
for _, _, channel, genre in drama_data:
    channel_count[channel] = channel_count.get(channel, 0) + 1
    genre_count[genre] = genre_count.get(genre, 0) + 1

doc.add_paragraph('频道分布：')
for ch, cnt in sorted(channel_count.items(), key=lambda x: -x[1]):
    doc.add_paragraph(f'  {ch}：{cnt}部（{cnt/len(drama_data)*100:.0f}%）')

doc.add_paragraph()
doc.add_paragraph('题材分布：')
for genre, cnt in sorted(genre_count.items(), key=lambda x: -x[1]):
    doc.add_paragraph(f'  {genre}：{cnt}部（{cnt/len(drama_data)*100:.0f}%）')

doc.add_paragraph()
doc.add_paragraph(f'总计：{len(drama_data)}部剧')

# Save
output_path = r'D:\mazal-folder\projects\data-analysis\剧目分类标签表.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
