# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys
sys.stdout.reconfigure(encoding='utf-8')

def set_run_font(run, font_name='楷体', font_size=12, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(doc, text, level=0):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_pic_center(doc, path, width=Inches(5.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
    return table

BASE = r'D:\mazal-folder\projects\data-analysis\assets\charts\4_南瓜'
OUT = r'D:\mazal-folder\projects\data-analysis\publish_erchuang_report.docx'

doc = Document()

# Title
add_title(doc, '第二部分：上架趋势与原创二创对比', level=0)
doc.add_paragraph()

# ===================== PART 1: 上架趋势 =====================
add_title(doc, '一、上架趋势', level=1)

add_body(doc, '4月期间（4/1-4/30），works_list 共收录发布剧目34部。上架节奏呈现"先抑后扬"特征：W1（4/1-5）无新剧上线，W2（清明周）起快速放量，W4（4/20-26）达到峰值11部，W5收尾9部。整体来看，4月中旬至月底是发布高峰期。')

add_pic_center(doc, f'{BASE}/publish_daily.png', width=Inches(5.5))
doc.add_paragraph()

add_body(doc, '图1为日频折线图，可见4月7日、14日、22日分别出现三个明显峰值，对应W2、W3、W4周期的启动日。整体发布节奏在工作日后半段集中，周初相对平淡。')

add_pic_center(doc, f'{BASE}/publish_weekly.png', width=Inches(4.5))
doc.add_paragraph()

add_body(doc, '图2为周频柱状图。W4（4/20-26）发布量最高，达11部，占全月32%；其次为W5的9部（26%）和W2的8部（24%）。W1无新剧上线。')

doc.add_paragraph()

# ===================== PART 2: 原创 vs 二创 =====================
add_title(doc, '二、原创账号 vs 二创账号对比', level=1)

add_body(doc, '以主账号"南瓜动漫短剧场"（原创）和"爱动漫微剧场"（二创）为对象，'
              '对两部交叉剧目"祭品公主的生存法则"与"替嫁后，沉睡老公超会撩"进行播放量、广告收益和万播收益对比。')

# Summary table
add_table(doc,
    ['账号', '剧目', '播放量', '广告收益（元）', '万播收益（元/万播）'],
    [
        ['南瓜动漫短剧场（原创）', '祭品公主的生存法则', '11,026', '0.06', '0.05'],
        ['爱动漫微剧场（二创）', '祭品公主的生存法则', '2,583', '0.00', '0.00'],
        ['南瓜动漫短剧场（原创）', '替嫁后，沉睡老公超会撩', '120,532', '0.80', '0.07'],
        ['爱动漫微剧场（二创）', '替嫁后，沉睡老公超会撩', '107,152', '0.46', '0.04'],
        ['南瓜动漫短剧场（合计）', '两剧合计', '131,558', '0.86', '0.07'],
        ['爱动漫微剧场（合计）', '两剧合计', '109,735', '0.46', '0.04'],
    ]
)
doc.add_paragraph()

add_pic_center(doc, f'{BASE}/orig_vs_erc.png', width=Inches(5.5))
doc.add_paragraph()

add_body(doc, '图表说明：图左为万播收益对比，图右为广告收益对比。'
              '两部剧中，原创账号"南瓜动漫短剧场"在播放量、收益和万播收益三个指标上均优于二创账号"爱动漫微剧场"。'
              '替嫁后一剧，两账号播放量相近（12.0万 vs 10.7万），但原创账号万播收益（0.07元）高于二创（0.04元），'
              '说明同一内容在原创账号下的变现效率更高。祭品公主一剧因整体量级较小，二创账号几乎没有收益产出。'
              '总体来看，原创账号加总万播收益0.07元/万播，二创账号0.04元/万播，原创账号效率约为二创的1.75倍。')

doc.add_paragraph()

# Save
doc.save(OUT)
print(f'Saved: {OUT}')