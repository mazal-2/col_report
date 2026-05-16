# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys
sys.stdout.reconfigure(encoding='utf-8')

OUT = r'D:\mazal-folder\projects\data-analysis\南瓜_4月运营分析综合报告.docx'

doc = Document()

# ===================== PAGE SETUP =====================
# A4: 210mm x 297mm
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ===================== STYLE HELPERS =====================
def make_paragraph_format(base_font_size, first_line_indent=0, line_spacing=22,
                           space_before=0, space_after=0,
                           alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                           bold=False):
    pf = doc.add_paragraph().paragraph_format
    pf.line_spacing = Pt(line_spacing)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.alignment = alignment
    if first_line_indent > 0:
        pf.first_line_indent = Pt(first_line_indent)
    return pf

def add_para(text, font_size, bold=False, indent=0, space_before=0, space_after=0,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=22, color=(0,0,0)):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(line_spacing)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = alignment
    if indent > 0:
        p.paragraph_format.first_line_indent = Pt(indent)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color)
    return p

def add_pic(path, width=Inches(5.0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)

def add_pic_inline(text, path, font_size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.add_run().add_picture(path, width=Inches(4.5))

# Inches helper for pic width
from docx.shared import Inches

# ===================== TITLE PAGE =====================
add_para('南瓜视频号 4月运营分析报告', 16, bold=True,
         space_before=6, space_after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('数据周期：2026年4月1日 — 4月30日', 11,
         space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ===================== 内部分析 =====================
add_para('内部分析', 15, bold=True,
         space_before=6, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para('第一部分：内容倾向', 12, bold=True,
         space_before=4, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)

BASE = r'D:\mazal-folder\projects\data-analysis\assets\charts\4_南瓜'

# ---- 1.1 日频 ----
add_para('（一）日频趋势', 11, bold=True,
         space_before=4, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para('以下展示4月（4/1-4/30）每日广告收益与万播收益的趋势，按男女频和题材两种维度分类。可观察到流量脉冲集中在中下旬，万播收益波动反映变现效率的日间变化。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_pic(f'{BASE}/chart1_daily_channel.png', width=Inches(4.8))
add_para('图 1  日频 — 男女频广告收益与万播收益趋势', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('女频因基数大主导整体趋势，男频万播收益更高说明变现效率更优。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_pic(f'{BASE}/chart2_daily_genre.png', width=Inches(4.8))
add_para('图 2  日频 — 题材广告收益与万播收益趋势', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('都市内容占据主导，但玄幻题材万播收益显著更高，说明小流量也能有高变现。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# ---- 1.2 周频 ----
add_para('（二）周频趋势', 11, bold=True,
         space_before=4, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para('将4月划分为5个自然周（W1:4/1-5, W2:4/6-12, W3:4/13-19, W4:4/20-26, W5:4/27-30），聚合每周的广告收益与万播收益。整体呈先升后降再反弹的走势，W3因清明小长假后复工达到峰值。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_pic(f'{BASE}/chart3_weekly_channel.png', width=Inches(4.5))
add_para('图 3  周频 — 男女频广告收益与万播收益', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('女频在W3达到峰值，男频在W4小幅回落但效率更稳定。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_pic(f'{BASE}/chart4_weekly_genre.png', width=Inches(4.5))
add_para('图 4  周频 — 题材广告收益与万播收益', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('都市内容W3爆发后回落（W4最低），玄幻在W4反而逆势上扬，说明题材间存在互补效应。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# ---- 1.3 月频 ----
add_para('（三）月频占比', 11, bold=True,
         space_before=4, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para('4月全月汇总，广告收益与阅读量的占比分布。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_pic(f'{BASE}/chart5_monthly_channel.png', width=Inches(4.5))
add_para('图 5  月频 — 男女频广告收益与阅读量占比', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('女频收益1651.39元（占比90.6%，阅读量93.7%），仍是绝对主力；男频收益170.87元（占比9.4%，阅读量6.3%），万播收益达25.95元/万播，远超女频的16.74元/万播，男频用户变现效率显著更优。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_pic(f'{BASE}/chart6_monthly_genre.png', width=Inches(4.5))
add_para('图 6  月频 — 题材广告收益与阅读量占比', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('都市收益1622.51元（占89.1%），玄幻177.09元（9.7%），末世22.66元（1.2%）；但万播收益方面，玄幻24.13元>都市17.43元>末世4.71元，与收益规模呈反向关系，说明小众题材反而变现效率更高。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# ---- 1.4 特定样本 ----
add_para('（四）特定样本', 11, bold=True,
         space_before=4, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_para('1，按剧目划分：分别选取广告收益、阅读量、万播收益三项指标各前3名。', 10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# Table data
table = doc.add_table(rows=4, cols=7)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['指标', '剧名', '广告收益（元）', '阅读量', '万播收益（元/万播）', '频道', '题材']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(0,0,0)

data_rows = [
    ['广告收益 Top1', '新婚后，大叔全家爆宠我', '848.49', '170,120', '49.88', '女频', '都市'],
    ['广告收益 Top2', '金牌替身', '310.42', '399,175', '7.78', '女频', '都市'],
    ['广告收益 Top3', '替嫁后，沉睡老公超会撩', '161.83', '243,744', '6.64', '女频', '都市'],
]
for ri, row_data in enumerate(data_rows):
    for ci, val in enumerate(row_data):
        cell = table.rows[ri+1].cells[ci]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0,0,0)

doc.add_paragraph()

table2 = doc.add_table(rows=4, cols=7)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(0,0,0)

data_rows2 = [
    ['阅读量 Top1', '金牌替身', '310.42', '399,175', '7.78', '女频', '都市'],
    ['阅读量 Top2', '替嫁后，沉睡老公超会撩', '161.83', '243,744', '6.64', '女频', '都市'],
    ['阅读量 Top3', '新婚后，大叔全家爆宠我', '848.49', '170,120', '49.88', '女频', '都市'],
]
for ri, row_data in enumerate(data_rows2):
    for ci, val in enumerate(row_data):
        cell = table2.rows[ri+1].cells[ci]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0,0,0)

doc.add_paragraph()

table3 = doc.add_table(rows=4, cols=7)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(0,0,0)

data_rows3 = [
    ['万播收益 Top1', '渊狐记：重生不负心', '94.98', '6,683', '142.12', '男频', '玄幻'],
    ['万播收益 Top2', '双鲤换生，这一世我所向披靡', '37.64', '3,266', '115.25', '男频', '玄幻'],
    ['万播收益 Top3', '极热65度：老公中奖2亿...', '59.10', '7,059', '83.72', '女频', '都市'],
]
for ri, row_data in enumerate(data_rows3):
    for ci, val in enumerate(row_data):
        cell = table3.rows[ri+1].cells[ci]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0,0,0)

doc.add_paragraph()

add_para('2，关键样本描述', 10, bold=True, space_before=4, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para('广告收益最高：新婚后，大叔全家爆宠我，总收益848.49元，总阅读量170,120，万播收益49.88元/万播。该剧为女频都市题材，4月中旬（4/15前后）开始起量，随后两周持续放量，成为4月最大爆款。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('阅读量最高：金牌替身，总阅读量399,175（全月最高），但总收益仅310.42元，万播收益7.78元/万播，变现效率仅为新婚后（49.88元）的1/6，存在高流量低变现问题。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('万播收益最高：渊狐记：重生不负心（男频玄幻），总收益94.98元，总阅读量6,683，万播收益高达142.12元/万播，是"小流量高变现"的典型代表。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# ===================== 第二部分：受众情况 =====================
add_para('第二部分：受众情况', 12, bold=True,
         space_before=6, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_para('（一）年龄分布', 11, bold=True,
         space_before=4, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para('本次用户年龄共划分七个档次，各档次对应绝对值依次为：83、50、94、342、213、101、27。用户年龄分布呈现中间高、两端低的纺锤形特征，用户集中度差异显著。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('七个年龄档次中，第四年龄档人数达到峰值342，为全年龄段用户体量最高梯队；其次为第五年龄档213人、第三年龄档94人；第一、第二、第六年龄档分别为83人、50人、101人；第七年龄档用户体量最少，仅27人。整体来看，用户核心群体高度集中在第四、第五中间年龄档次，低龄段与高龄段用户体量均明显偏少。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_pic(f'{BASE}/用户受众_年龄.png', width=Inches(4.5))
add_para('图 7  用户年龄分布柱状图', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('可直观看到七个年龄档次的人数落差，中间档位形成明显高峰，首尾档位用户规模持续走低，契合整体纺锤形分布规律。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_para('（二）性别分布', 11, bold=True,
         space_before=4, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para('用户性别划分为男性、女性、未知三个档次，对应体量数据依次为169、643、98。整体性别结构分化特征十分突出。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('从各档次数据来看，女性用户体量遥遥领先，达643人，是平台用户的绝对主力群体；男性用户规模为169人，远低于女性用户体量；性别标注未知用户仅有98人，在三大档次中体量最小。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_pic(f'{BASE}/用户受众_性别.png', width=Inches(4.5))
add_para('图 8  用户性别分布对比图', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('清晰呈现出女性用户占比远超男性与未知群体的格局，整体用户性别呈现女性主导、男性补充、未知占比偏低的分布特点。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# ===================== 第三部分：运营方式 =====================
add_para('第三部分：运营方式', 12, bold=True,
         space_before=6, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)

# 3.1 上架节奏
add_para('3.1 上架节奏', 11, bold=True,
         space_before=4, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para('4月期间（4/1-4/30），works_list共收录发布剧目34部。上架节奏呈现"先抑后扬"特征：W1（4/1-5）无新剧上线，W2（清明周）起快速放量，W4（4/20-26）达到峰值11部，W5收尾9部。整体来看，4月中旬至月底是发布高峰期。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_pic(f'{BASE}/publish_daily.png', width=Inches(4.5))
add_para('图 9  发布剧数日频折线图', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('可见4月7日、14日、22日分别出现三个明显峰值，对应W2、W3、W4周期的启动日。整体发布节奏在工作日后半段集中，周初相对平淡。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_pic(f'{BASE}/publish_weekly.png', width=Inches(4.0))
add_para('图 10  发布剧数周频柱状图', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('W4（4/20-26）发布量最高，达11部，占全月32%；其次为W5的9部（26%）和W2的8部（24%）。W1无新剧上线。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# 3.2 二创
add_para('3.2 原创账号与二创账号对比', 11, bold=True,
         space_before=4, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_para('以主账号"南瓜动漫短剧场"（原创）和"爱动漫微剧场"（二创）为对象，对两部交叉剧目"祭品公主的生存法则"与"替嫁后，沉睡老公超会撩"进行播放量、广告收益和万播收益对比。',
         10, indent=21, space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# Erchuang table
table4 = doc.add_table(rows=7, cols=5)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['账号', '剧目', '播放量', '广告收益（元）', '万播收益（元/万播）']
for i, h in enumerate(headers2):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(0,0,0)

erc_data = [
    ['南瓜动漫短剧场（原创）', '祭品公主的生存法则', '11,026', '6.04', '5.48'],
    ['爱动漫微剧场（二创）', '祭品公主的生存法则', '2,583', '0.05', '0.19'],
    ['南瓜动漫短剧场（原创）', '替嫁后，沉睡老公超会撩', '120,532', '79.51', '6.60'],
    ['爱动漫微剧场（二创）', '替嫁后，沉睡老公超会撩', '107,152', '45.89', '4.28'],
    ['南瓜动漫短剧场（合计）', '两剧合计', '131,558', '85.55', '6.50'],
    ['爱动漫微剧场（合计）', '两剧合计', '109,735', '45.94', '4.19'],
]
for ri, row_data in enumerate(erc_data):
    for ci, val in enumerate(row_data):
        cell = table4.rows[ri+1].cells[ci]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0,0,0)

doc.add_paragraph()

add_pic(f'{BASE}/orig_vs_erc.png', width=Inches(5.0))
add_para('图 11  原创账号与二创账号万播收益及广告收益对比', 10, bold=True,
         space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('图左为万播收益对比，图右为广告收益对比。两部剧中，原创账号"南瓜动漫短剧场"在播放量、收益和万播收益三个指标上均优于二创账号"爱动漫微剧场"。替嫁后一剧，两账号播放量相近（12.0万 vs 10.7万），但原创账号万播收益（6.60元）高于二创（4.28元），说明同一内容在原创账号下的变现效率更高。总体来看，原创账号加总万播收益6.50元/万播，二创账号4.19元/万播，原创账号效率约为二创的1.55倍。',
         10, indent=21, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# ===================== SAVE =====================
doc.save(OUT)
print(f'Saved: {OUT}')