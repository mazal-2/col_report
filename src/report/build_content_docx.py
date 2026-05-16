# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd
import sys
sys.stdout.reconfigure(encoding='utf-8')

# ===================== DATA =====================
april_xl = r'D:\mazal-folder\projects\data-analysis\data\515_shipinhao\4_南瓜\april_daily_data.xlsx'
tag_csv = r'D:\mazal-folder\projects\data-analysis\data\515_shipinhao\4_南瓜\剧目分类标签_full.csv'
xl = pd.ExcelFile(april_xl)
df_sum = xl.parse('剧名汇总')
tag_df = pd.read_csv(tag_csv)
monthly = df_sum.merge(tag_df, on='剧名', how='left')
monthly['万播收益'] = monthly['广告收益'] / monthly['阅读量'] * 10000
monthly['万播收益'] = monthly['万播收益'].replace([float('inf'), float('-inf')], 0).fillna(0)

monthly_channel = monthly.groupby('频道').agg(
    广告收益=('广告收益', 'sum'),
    阅读量=('阅读量', 'sum')
).reset_index()
monthly_channel['万播收益'] = monthly_channel['广告收益'] / monthly_channel['阅读量'] * 10000

monthly_genre = monthly.groupby('题材').agg(
    广告收益=('广告收益', 'sum'),
    阅读量=('阅读量', 'sum')
).reset_index()
monthly_genre['万播收益'] = monthly_genre['广告收益'] / monthly_genre['阅读量'] * 10000

# ===================== DOCX =====================
doc = Document()

# Title style helper
def set_run_font(run, font_name='楷体', font_size=12, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(doc, text, level=0):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_pic_center(doc, path, width=Inches(5.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)

def add_table(doc, headers, rows, header_color=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    # Rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
    return table

BASE = r'D:\mazal-folder\projects\data-analysis\assets\charts\4_南瓜'

# ===================== TITLE =====================
add_title(doc, '第一部分：内容倾向分析', level=0)
doc.add_paragraph()

# ===================== SECTION 1: 日频 =====================
add_title(doc, '一、日频趋势（折线图）', level=1)

add_body(doc, '以下展示4月（4/1-4/30）每日广告收益与万播收益的趋势，按男女频和题材两种维度分类。'
              '可观察到流量脉冲集中在中下旬，万播收益波动反映变现效率的日间变化。')

add_pic_center(doc, f'{BASE}/chart1_daily_channel.png', width=Inches(5.5))
doc.add_paragraph()

add_body(doc, '图1-1为男女频日频收益对比，图1-2为万播收益对比。'
              '女频因基数大主导整体趋势，男频万播收益更高说明变现效率更优。')

add_pic_center(doc, f'{BASE}/chart2_daily_genre.png', width=Inches(5.5))
doc.add_paragraph()

add_body(doc, '图2-1为题材维度日频收益对比，图2-2为万播收益。'
              '都市内容占据主导，但玄幻题材万播收益显著更高，说明小流量也能有高变现。')

doc.add_paragraph()

# ===================== SECTION 2: 周频 =====================
add_title(doc, '二、周频趋势（柱形图）', level=1)

add_body(doc, '将4月划分为5个自然周（W1:4/1-5, W2:4/6-12, W3:4/13-19, W4:4/20-26, W5:4/27-30），'
              '聚合每周的广告收益与万播收益。整体呈先升后降再反弹的走势，W3因清明小长假后复工达到峰值。')

add_pic_center(doc, f'{BASE}/chart3_weekly_channel.png', width=Inches(5.0))
doc.add_paragraph()

add_body(doc, '图3为男女频周频对比。女频在W3达到峰值，男频在W4小幅回落但效率更稳定。')

add_pic_center(doc, f'{BASE}/chart4_weekly_genre.png', width=Inches(5.0))
doc.add_paragraph()

add_body(doc, '图4为题材周频对比。都市内容W3爆发后回落（W4最低），玄幻在W4反而逆势上扬，说明题材间存在互补效应。')

doc.add_paragraph()

# ===================== SECTION 3: 月频 =====================
add_title(doc, '三、月频占比（饼图）', level=1)

add_body(doc, '4月全月汇总，广告收益与阅读量的占比分布。')

add_pic_center(doc, f'{BASE}/chart5_monthly_channel.png', width=Inches(5.5))
doc.add_paragraph()

add_body(doc, '图5：男女频月频占比。女频收益1651.39元（占比90.6%，阅读量93.7%），仍是绝对主力；'
              '男频收益170.87元（占比9.4%，阅读量6.3%），万播收益达25.95元/万播，远超女频的16.74元/万播，'
              '说明男频用户变现效率显著更优。')

add_pic_center(doc, f'{BASE}/chart6_monthly_genre.png', width=Inches(5.5))
doc.add_paragraph()

add_body(doc, '图6：题材月频占比。都市收益1622.51元（占89.1%），玄幻177.09元（9.7%），末世22.66元（1.2%）；'
              '但万播收益方面，玄幻24.13元>都市17.43元>末世4.71元，与收益规模呈反向关系，说明小众题材反而变现效率更高。')

doc.add_paragraph()

# ===================== SECTION 4: 特定样本 =====================
add_title(doc, '四、特定样本分析', level=1)

add_body(doc, '1，按剧目划分：分别选取广告收益、阅读量、万播收益三项指标各前3名')

# Revenue top3
rev_top3 = monthly.nlargest(3, '广告收益')[['剧名', '广告收益', '阅读量', '万播收益', '频道', '题材']]
rev_top3['万播收益'] = rev_top3['万播收益'].round(2)
add_table(doc,
    ['指标', '剧名', '广告收益（元）', '阅读量', '万播收益（元/万播）', '频道', '题材'],
    [[f'广告收益 Top{i+1}', r['剧名'], f"{round(r['广告收益'], 2):,.2f}", f"{int(r['阅读量']):,}", round(r['万播收益'], 2), r['频道'], r['题材']]
     for i, (_, r) in enumerate(rev_top3.iterrows())]
)
doc.add_paragraph()

# Read top3
read_top3 = monthly.nlargest(3, '阅读量')[['剧名', '广告收益', '阅读量', '万播收益', '频道', '题材']]
add_table(doc,
    ['指标', '剧名', '广告收益（元）', '阅读量', '万播收益（元/万播）', '频道', '题材'],
    [[f'阅读量 Top{i+1}', r['剧名'], f"{round(r['广告收益'], 2):,.2f}", f"{int(r['阅读量']):,}", round(r['万播收益'], 2), r['频道'], r['题材']]
     for i, (_, r) in enumerate(read_top3.iterrows())]
)
doc.add_paragraph()

# Wanbo top3
wanbo_top3 = monthly[monthly['阅读量'] > 0].nlargest(3, '万播收益')[['剧名', '广告收益', '阅读量', '万播收益', '频道', '题材']]
add_table(doc,
    ['指标', '剧名', '广告收益（元）', '阅读量', '万播收益（元/万播）', '频道', '题材'],
    [[f'万播收益 Top{i+1}', r['剧名'], f"{round(r['广告收益'], 2):,.2f}", f"{int(r['阅读量']):,}", round(r['万播收益'], 2), r['频道'], r['题材']]
     for i, (_, r) in enumerate(wanbo_top3.iterrows())]
)
doc.add_paragraph()

add_body(doc, '2，关键样本描述')
add_body(doc, '广告收益最高：新婚后，大叔全家爆宠我，总收益848.49元，总阅读量170,120，万播收益49.88元/万播。'
              '该剧为女频都市题材，4月中旬（4/15前后）开始起量，随后两周持续放量，成为4月最大爆款。')
add_body(doc, '阅读量最高：金牌替身，总阅读量399,175（全月最高），但总收益仅310.42元，万播收益7.78元/万播，'
              '变现效率仅为新婚后（49.88元）的1/6，存在高流量低变现问题。')
add_body(doc, '万播收益最高：渊狐记：重生不负心（男频玄幻），总收益94.98元，总阅读量6,683，万播收益高达142.12元/万播，'
              '是"小流量高变现"的典型代表。')

doc.add_paragraph()

# ===================== SAVE =====================
output = r'D:\mazal-folder\projects\data-analysis\content_trend_report.docx'
doc.save(output)
print(f'Saved: {output}')